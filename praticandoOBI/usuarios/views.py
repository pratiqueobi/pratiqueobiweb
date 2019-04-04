# -*- coding: utf-8 -*-
import codecs
from django.shortcuts import render, redirect, get_object_or_404
import os
from .forms import ProfileForm, ProvaForm, QuestoesForm
from provasobi.models import ProvaPerson, Prova, Questao, Classificacao, Problema, Alternativa
from django.contrib import messages
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from django.core.files.storage import FileSystemStorage
from django.http import HttpResponse

from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from django.shortcuts import render, redirect
from django.contrib.auth import login, authenticate
from django.contrib.sites.shortcuts import get_current_site
from django.utils.encoding import force_bytes, force_text
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from django.template.loader import render_to_string
from .tokens import account_activation_token
from django.contrib.auth.models import User
from django.core.mail import EmailMessage

# API DRIVE:
from googleapiclient.discovery import build
from httplib2 import Http
from oauth2client import file, client, tools
from apiclient.http import MediaFileUpload

from watson import search as watson

# If modifying these scopes, delete the file token.json.
SCOPES = 'https://www.googleapis.com/auth/drive'


# @login_required
def home_usuario(request):
    return render(request, 'usuarios/homeusuario.html', {})


def update_perfil(request):
    if request.method == 'POST':
        user_form = UserForm(request.POST, instance=request.user)
        profile_form = ProfileForm(request.POST, instance=request.user.profile)
        if user_form.is_valid() and profile_form.is_valid():
            user_form.save()
            profile_form.save()
            messages.success(request, 'Perfil atualizado.')
            return redirect('home')
        else:
            messages.error(request, 'Corrija os erros.')
    else:
        user_form = UserForm(instance=request.user)
        profile_form = ProfileForm(instance=request.user.profile)
    return render(request, 'usuarios/perfil.html', {'user_form': user_form, 'profile_form': profile_form})


def cadastro_perfil(request):
    error = False
    if request.method == 'POST':
        form = ProfileForm(request.POST)
        if form.is_valid():
            error = False
            user = form.save()
            user.refresh_from_db()  # load the profile instance created by the signal
            user.profile.data_nascimento = form.cleaned_data.get('data_nascimento')
            print(user.profile.data_nascimento)
            user.profile.instituicao = form.cleaned_data.get('instituicao')
            user.profile.localizacao = form.cleaned_data.get('localizacao')

            to_email = form.cleaned_data.get('email')
            user.is_active = False
            user.save()

            current_site = get_current_site(request)
            print(current_site)  # AQUI: SETAR SITE PRO HEROKU
            mail_subject = 'Ative sua conta no Praticando OBI'
            message = render_to_string('registration/acc_active_email.html', {
                'user': user,
                'domain': current_site.domain,
                'uid': urlsafe_base64_encode(force_bytes(user.pk)).decode(),
                'token': account_activation_token.make_token(user),
            })

            email = EmailMessage(mail_subject, message, to=[to_email])
            email.send()
            return render(request, 'registration/confirmacao.html', {})
        else:
            error = True
    else:
        form = ProfileForm()
    return render(request, 'usuarios/signup.html', {'form': form, 'error': error})


def activate(request, uidb64, token):
    try:
        uid = force_text(urlsafe_base64_decode(uidb64))
        user = User.objects.get(pk=uid)
    except(TypeError, ValueError, OverflowError, User.DoesNotExist):
        user = None
    if user is not None and account_activation_token.check_token(user, token):
        user.is_active = True
        user.save()
        login(request, user)
        # return redirect('home')
        messages.success(request, 'Obrigado por confirmar sua conta!')
        return render(request, 'registration/login.html', {})
    else:
        return redirect('home')


def provaperson(request):
    if request.method == "POST":
        form = ProvaForm(request.POST)
        if form.is_valid():
            provaperson = form.save(commit=False)
            provaperson.autor = request.user.profile
            provaperson.save()
            return redirect('usuarios_obi:questoes_busca', provaperson.pk)
    else:
        form = ProvaForm()
    return render(request, 'novasprovas/provaperson.html', {'form': form})


def provaperson_excluir(request, pk):
    provaperson = get_object_or_404(ProvaPerson, pk=pk)
    provaperson.delete()
    return redirect('usuarios_obi:provasperson')


def provaperson_edit(request, pk):
    provaperson = get_object_or_404(ProvaPerson, pk=pk)
    if request.method == "POST":
        form = ProvaForm(request.POST, instance=provaperson)
        if form.is_valid():
            provaperson = form.save(commit=False)
            provaperson.autor = request.user.profile
            provaperson.save()
            messages.success(request, "Alterações salvas!")
            return redirect('usuarios_obi:provaperson_edit', provaperson.pk)
    else:
        form = ProvaForm(instance=provaperson)

    return render(request, 'novasprovas/provaperson_edit.html',
                  {'form': form, 'pk': pk, 'titulo': provaperson.titulo, 'ano': provaperson.ano,
                   'dificuldade': provaperson.dificuldade, 'obs': provaperson.observacoes})


# mostra as provas criadas
def provasperson(request):
    provas = ProvaPerson.objects.filter(autor=request.user.profile)
    return render(request, 'minhasprovas.html', {'provas': provas})


def provaperson_detail(request, pk):
    provaperson = ProvaPerson.objects.all().filter(pk=pk)
    return render(request, 'novasprovas/provaperson_detail.html', {'provaperson': provaperson})


def gabarito(request, codprova):
    provaperson = get_object_or_404(ProvaPerson, pk=codprova, autor=request.user.profile)
    questoes = Questao.objects.all().filter(codquestao__in=provaperson.questoes.all()).order_by('numeroquestao')

    id_problemas = Questao.objects.all().filter(codquestao__in=provaperson.questoes.all()).values('codproblema')
    problemas = Problema.objects.all().filter(codproblema__in=id_problemas).distinct()

    return render(request, 'novasprovas/gabarito.html',
                  {'provaperson': provaperson, 'problemas': problemas, 'questoes': questoes, 'codprova': codprova})


def busca(request):
    error = False
    provas = Prova.objects.all()
    q = request.GET.get('q', '')
    problemas = Problema.objects.all()
    if q:
        problemas = watson.filter(problemas, q)

    return render(request, 'novasprovas/addquestoes.html', {'provas': provas,
                                                            'error': error,
                                                            'problemas': problemas})

def questoes_busca(request, pk):
    provaperson = get_object_or_404(ProvaPerson, pk=pk, autor=request.user.profile)
    q = request.GET.get('q', '')
    problemas = Problema.objects.all()
    if q:
        problemas = watson.filter(problemas, q)

    if request.method == 'POST':
        id_questoes = request.POST.getlist('checks')

        for q in id_questoes:
            provaperson.questoes.add(q)

        if request.POST.get('finalizar', False):
            return redirect('usuarios_obi:provaperson_pronta', codprova=pk)

    return render(request, 'novasprovas/addquestoes.html',
                  {'provaperson': provaperson, 'pk': pk, 'problemas': problemas})


def questoes_add(request, codproblema, pk):
    problema = Problema.objects.get(codproblema=codproblema)
    questoes = Questao.objects.filter(codproblema=codproblema).order_by('numeroquestao')  # .filter(codproblema__in=id_questoes)
    return render(request, 'novasprovas/addquestoes_select.html',
                  {'problema': problema, 'questoes': questoes, 'pk': pk})


def provaperson_pronta(request, codprova):
    provaperson = get_object_or_404(ProvaPerson, pk=codprova, autor=request.user.profile)
    questoes = Questao.objects.all().filter(codquestao__in=provaperson.questoes.all()).order_by('numeroquestao')

    id_problemas = Questao.objects.all().filter(codquestao__in=provaperson.questoes.all()).values('codproblema')
    problemas = Problema.objects.all().filter(codproblema__in=id_problemas).distinct()

    return render(request, 'novasprovas/provaperson_pronta.html',
                  {'provaperson': provaperson, 'problemas': problemas, 'questoes': questoes, 'codprova': codprova})


acentoserro = ["á", "à", "ã", "Ã", "é", "ê", "õ", "ô", "ó", "ç", "ú", "ı́"]
acentos = ["á", "à", "ã", "Ã", "é", "ê", "õ", "ô", "ó", "ç", "ú", "í"]

def provaperson_baixar(request, codprova):

    provaperson = get_object_or_404(ProvaPerson, pk=codprova, autor=request.user.profile)
    questoes = Questao.objects.all().filter(codquestao__in=provaperson.questoes.all()).order_by('numeroquestao')

    id_problemas = Questao.objects.all().filter(codquestao__in=provaperson.questoes.all()).values('codproblema')
    problemas = Problema.objects.all().filter(codproblema__in=id_problemas).distinct()

    count = 1
    doc = SimpleDocTemplate("/tmp/prova-" + str(codprova) + ".pdf", rightMargin=50, leftMargin=50, topMargin=40, bottomMargin=50)
    styles = getSampleStyleSheet()
    Story = [Spacer(1, 0.2 * inch)]
    style = styles["Normal"]
    par = Paragraph('<para align=center fontSize=20 > <b>' + provaperson.titulo + '</b></para>', style)
    Story.append(par)
    Story.append(Spacer(1, 0.4 * inch))

    for p in problemas:
        par = Paragraph('<para align=center fontSize=14><b>' + p.tituloproblema + '</b><br/></para>', style)
        Story.append(par)
        Story.append(Spacer(1, 0.2 * inch))

        e = p.enunciadoproblema
        print(e)
        for i in range(len(acentos)):
            e = e.replace(acentoserro[i], acentos[i])

        par = Paragraph('<para fontSize=12>' + e + '<br/></para>', style)
        Story.append(par)

        if p.regrasproblema:
            Story.append(Spacer(1, 0.1 * inch))
            par = Paragraph('<para fontSize=12><b>REGRAS:</b><br/></para>', style)
            Story.append(par)

            e = p.regrasproblema
            print(e)
            for i in range(len(acentos)):
                e = e.replace(acentoserro[i], acentos[i])

            par = Paragraph('<para fontSize=12>' + e + '<br/></para>', style)
            Story.append(par)

        Story.append(Spacer(1, 0.2 * inch))

        if p.imgproblema:
            # local: 'static/'
            # heroku: '/app/praticandoOBI/static/'
            img = Image('/app/praticandoOBI/static/' + p.imgproblema, 4 * inch, 4 * inch)
            Story.append(img)

        for q in questoes:
            if p.codproblema == q.codproblema.codproblema:
                e = q.enunciadoquestao
                for i in range(len(acentos)):
                    e = e.replace(acentoserro[i], acentos[i])

                par = Paragraph('<para fontSize=12><b>Questão ' + str(count) + "</b> - " + e + '<br/></para>', style)
                Story.append(par)
                Story.append(Spacer(1, 0.2 * inch))
                count += 1

                if q.imgquestao:
                    # local: 'static/'
                    # heroku: '/app/praticandoOBI/static/'
                    img = Image('/app/praticandoOBI/static/' + q.imgproblema, 2 * inch, 2 * inch)
                    Story.append(img)

                for a in q.get_alternativas():
                    e = a['textoalternativa']
                    for i in range(len(acentos)):
                        e = e.replace(acentoserro[i], acentos[i])

                    par = Paragraph('<para fontSize=12><b>' + a['letraalternativa'] + ')</b> ' + e + '<br/></para>',
                                    style)
                    Story.append(par)
                    Story.append(Spacer(1, 0.1 * inch))

        Story.append(Spacer(1, 0.3 * inch))
    doc.build(Story)
    nome = "prova-" + str(codprova)
    fs = FileSystemStorage("/tmp")
    with fs.open("prova-" + str(codprova) + ".pdf") as pdf:
        response = HttpResponse(pdf, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename=' + nome + '.pdf'
    return response


def provaperson_baixar_docx(request, codprova):

    # ENCONTRA CONTEUDO DA PROVA:
    provaperson = get_object_or_404(ProvaPerson, pk=codprova, autor=request.user.profile)
    questoes = Questao.objects.all().filter(codquestao__in=provaperson.questoes.all()).order_by('numeroquestao')

    id_questoes = []
    for q in questoes:
        # print(q.enunciadoquestao) #ARRUMAR ACENTUAÇÃO
        id_questoes.append(q)

    alternativas = Alternativa.objects.all().select_related('codquestao').filter(codquestao__in=id_questoes)

    id_problemas = Questao.objects.all().filter(codquestao__in=provaperson.questoes.all()).values('codproblema')
    problemas = Problema.objects.all().filter(codproblema__in=id_problemas).distinct()


    document = Document()
    document.add_heading(provaperson.titulo, 0)

    count = 1
    for p in problemas:
        par2 = document.add_heading(p.tituloproblema, level=1)
        par2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = par2.add_run()

        run2.add_break()

        par1 = document.add_paragraph()
        # par1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = par1.add_run(p.enunciadoproblema)
        run.add_break()
        run.add_break()

        if p.regrasproblema:
            par1.add_run('REGRAS: ')
            run = par1.add_run(p.regrasproblema)
            run.add_break()
            run.add_break()

        if p.imgproblema:
            # local: 'static/ + p.imgproblema'
            # heroku: '/app/praticandoOBI/static/'
            document.add_picture('/app/praticandoOBI/static/' + p.imgproblema, width=Inches(4))

        par = document.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for q in questoes:
            if p.codproblema == q.codproblema.codproblema:

                e = q.enunciadoquestao
                for i in range(len(acentos)):
                    e = e.replace(acentoserro[i], acentos[i])

                par.add_run('Questão ').bold = True
                par.add_run(str(count)).bold = True
                par.add_run(': ')
                run = par.add_run(e)
                run.add_break()
                count += 1

                if q.imgquestao:
                    # local: 'static/ + q.imgproblema'
                    # heroku: '/app/praticandoOBI/static/' + q.imgproblema
                    document.add_picture('/app/praticandoOBI/static/' + q.imgproblema, width=Inches(4))

                for a in q.get_alternativas():
                    alt = a['textoalternativa']
                    for i in range(len(acentos)):
                        alt = alt.replace(acentoserro[i], acentos[i])

                    par.add_run(a['letraalternativa']).bold = True
                    par.add_run(') ').bold = True
                    run = par.add_run(alt)
                    run.add_break()

                run.add_break()
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=download.docx'
    document.save(response)

    return response


def upload_drive(request, codprova):
    # CRIA PROVA DOCX:
    provaperson = get_object_or_404(ProvaPerson, pk=codprova, autor=request.user.profile)
    questoes = Questao.objects.all().filter(codquestao__in=provaperson.questoes.all()).order_by('numeroquestao')

    id_problemas = Questao.objects.all().filter(codquestao__in=provaperson.questoes.all()).values('codproblema')
    problemas = Problema.objects.all().filter(codproblema__in=id_problemas).distinct()

    document = Document()
    document.add_heading(provaperson.titulo, 0)

    count = 1
    for p in problemas:
        par2 = document.add_heading(p.tituloproblema, level=1)
        par2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = par2.add_run()

        run2.add_break()

        par1 = document.add_paragraph()
        # par1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = par1.add_run(p.enunciadoproblema)
        run.add_break()
        run.add_break()

        if p.regrasproblema:
            par1.add_run('REGRAS: ')
            run = par1.add_run(p.regrasproblema)
            run.add_break()
            run.add_break()

        if p.imgproblema:
            # local: 'static/ + p.imgproblema'
            # heroku: '/app/praticandoOBI/static/'
            document.add_picture('/app/praticandoOBI/static/' + p.imgproblema, width=Inches(4))

        par = document.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for q in questoes:
            if p.codproblema == q.codproblema.codproblema:

                e = q.enunciadoquestao
                for i in range(len(acentos)):
                    e = e.replace(acentoserro[i], acentos[i])

                par.add_run('Questão ').bold = True
                par.add_run(str(count)).bold = True
                par.add_run(': ')
                run = par.add_run(e)
                run.add_break()
                count += 1

                if q.imgquestao:
                    #local: 'static/ + q.imgproblema'
                    #heroku: '/app/praticandoOBI/static/' + q.imgproblema
                    document.add_picture('/app/praticandoOBI/static/' + q.imgproblema, width=Inches(4))

                for a in q.get_alternativas():
                    alt = a['textoalternativa']
                    for i in range(len(acentos)):
                        alt = alt.replace(acentoserro[i], acentos[i])

                    par.add_run(a['letraalternativa']).bold = True
                    par.add_run(') ').bold = True
                    run = par.add_run(alt)
                    run.add_break()

                run.add_break()
    # SALVA A PROVA LOCAL
    document.save('/app/praticandoOBI/' + provaperson.titulo + '.docx')

    # CONECTA COM CONTA NO DRIVE
    store = file.Storage('/app/praticandoOBI/token.json')
    creds = store.get()

    if not creds or creds.invalid:
        flags = tools.argparser.parse_args(args=[])
        flow = client.flow_from_clientsecrets('/app/praticandoOBI/credentials.json', SCOPES)
        creds = tools.run_flow(flow, store, flags)
    service = build('drive', 'v3', http=creds.authorize(Http()))

    # ENVIA O ARQUIVO
    file_metadata = {
        'name': provaperson.titulo + '.docx',
    }
    media = MediaFileUpload('/app/praticandoOBI/' + provaperson.titulo + '.docx',
                            mimetype='text/',
                            resumable=True)
    prova = service.files().create(body=file_metadata,
                                   media_body=media).execute()

    # APAGA A PROVA SALVA LOCALMENTE
    os.remove('/app/praticandoOBI/' + provaperson.titulo + '.docx')

    return redirect('https://docs.google.com/document/d/' + prova.get('id') + '/edit')


def dadosbanco(request):
    file_path = '/app/praticandoOBI/OBI.db'
    with open(file_path, 'rb') as fh:
        response = HttpResponse(fh.read(), content_type="application/vnd.ms-excel")
        response['Content-Disposition'] = 'inline; filename="OBI.db"'
        return response
    raise Http404
